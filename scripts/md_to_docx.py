"""
MD → DOCX Dönüşüm Scripti

Selin Uyar dava klasöründeki tüm .md dosyalarını .docx'e çevirir.
- Frontmatter (YAML) otomatik atlanır
- Başlıklar (# ## ###) Word heading stiline dönüşür
- Kalın (**) ve italik (*) korunur
- Tablolar Word tablosuna dönüşür
- Liste öğeleri (-, 1-), *) korunur
- Code blokları (```) monospace font

Kullanım:
    python md_to_docx.py <kaynak_klasor> [hedef_klasor]

Örnek:
    python md_to_docx.py "G:/Drive'ım/Hukuk Bürosu/Aktif Davalar/Selin Uyar Kira Tespit Davası"
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def strip_frontmatter(text: str) -> str:
    """YAML frontmatter bloğunu (--- ... ---) kaldır."""
    if text.startswith('---'):
        end = text.find('---', 3)
        if end != -1:
            return text[end + 3:].lstrip('\n')
    return text


def strip_inline_format(text: str) -> list:
    """**kalın** ve *italik* ifadeleri parçala. [('düz metin', 'normal'), ('kalın', 'bold'), ...]"""
    parts = []
    pattern = re.compile(r'(\*\*[^\*]+\*\*|\*[^\*]+\*|`[^`]+`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], 'normal'))
        token = m.group()
        if token.startswith('**') and token.endswith('**'):
            parts.append((token[2:-2], 'bold'))
        elif token.startswith('`') and token.endswith('`'):
            parts.append((token[1:-1], 'code'))
        elif token.startswith('*') and token.endswith('*'):
            parts.append((token[1:-1], 'italic'))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], 'normal'))
    return parts


def add_formatted_text(paragraph, text: str):
    """Bir paragrafa inline formatlı metni ekler."""
    for chunk, style in strip_inline_format(text):
        run = paragraph.add_run(chunk)
        if style == 'bold':
            run.bold = True
        elif style == 'italic':
            run.italic = True
        elif style == 'code':
            run.font.name = 'Consolas'
            run.font.size = Pt(10)


def parse_table(lines: list, start: int) -> tuple:
    """Markdown tablosu ayrıştırır. (tablo_satırları, bitiş_indeks) döner."""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        if re.match(r'^\|[\s:|-]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i


def md_to_docx(md_path: Path, docx_path: Path):
    """MD dosyasını DOCX'e çevirir."""
    text = md_path.read_text(encoding='utf-8')
    text = strip_frontmatter(text)

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    lines = text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('```'):
            if in_code_block:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if stripped.startswith('|') and i + 1 < len(lines) and re.match(r'^\|[\s:|-]+\|$', lines[i + 1].strip()):
            rows, end = parse_table(lines, i)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'
                for r_idx, row_data in enumerate(rows):
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < len(table.rows[r_idx].cells):
                            cell = table.rows[r_idx].cells[c_idx]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            add_formatted_text(p, cell_text)
                            if r_idx == 0:
                                for run in p.runs:
                                    run.bold = True
            i = end
            continue

        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:].strip(), level=1)
            p.style.font.size = Pt(16)
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:].strip(), level=2)
            p.style.font.size = Pt(14)
        elif stripped.startswith('### '):
            p = doc.add_heading(stripped[4:].strip(), level=3)
            p.style.font.size = Pt(12)
        elif stripped.startswith('#### '):
            p = doc.add_heading(stripped[5:].strip(), level=4)
        elif re.match(r'^[-*+]\s+', stripped):
            content = re.sub(r'^[-*+]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, content)
        elif re.match(r'^\d+[.)]\s+', stripped):
            content = re.sub(r'^\d+[.)]\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, content)
        elif re.match(r'^\d+-\)\s+', stripped):
            p = doc.add_paragraph()
            add_formatted_text(p, stripped)
        elif stripped.startswith('---'):
            p = doc.add_paragraph('_' * 60)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith('> '):
            p = doc.add_paragraph(style='Intense Quote')
            add_formatted_text(p, stripped[2:])
        elif stripped == '':
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line)

        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    return docx_path


def convert_folder(src: Path, dst: Path = None):
    """Klasördeki tüm .md dosyalarını .docx'e çevirir."""
    if dst is None:
        dst = src

    md_files = list(src.rglob('*.md'))
    if not md_files:
        print(f'[UYARI] {src} klasöründe .md dosyası bulunamadı')
        return []

    print(f'[INFO] {len(md_files)} MD dosyası bulundu, DOCX dönüşümü başlıyor...')
    results = []
    for md_path in md_files:
        rel = md_path.relative_to(src)
        docx_path = dst / rel.with_suffix('.docx')
        try:
            md_to_docx(md_path, docx_path)
            print(f'  OK  {rel} -> {rel.with_suffix(".docx")}')
            results.append(docx_path)
        except Exception as e:
            print(f'  HATA {rel}: {e}')
    print(f'[INFO] Toplam {len(results)} DOCX üretildi.')
    return results


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Kullanim: python md_to_docx.py <kaynak_klasor> [hedef_klasor]')
        sys.exit(1)

    # Faz 1.4 profiling wrapper
    try:
        import os as _os
        sys.path.insert(0, _os.path.dirname(_os.path.abspath(__file__)))
        from _timing import TimedScript
        with TimedScript(__file__):
            src = Path(sys.argv[1])
            dst = Path(sys.argv[2]) if len(sys.argv) > 2 else None
            convert_folder(src, dst)
    except ImportError:
        src = Path(sys.argv[1])
        dst = Path(sys.argv[2]) if len(sys.argv) > 2 else None
        convert_folder(src, dst)
